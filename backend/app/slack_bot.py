"""
Slack Bot implementation using Slack Bolt
"""
import io
import logging
from typing import Any, Dict

import httpx
from app.config import settings
from app.database import Document, SessionLocal, SlackMessage
from app.rag import rag_engine
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
from slack_sdk import WebClient

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class SlackBot:
    """Wingman Slack Bot with RAG capabilities"""
    
    def __init__(self):
        """Initialize Slack Bot"""
        self.app = App(
            token=settings.SLACK_BOT_TOKEN,
            signing_secret=settings.SLACK_SIGNING_SECRET
        )
        self.client = WebClient(token=settings.SLACK_BOT_TOKEN)
        
        # Register event handlers
        self._register_handlers()
    
    def _register_handlers(self):
        """Register Slack event handlers"""
        
        @self.app.event("app_mention")
        def handle_mention(event, say):
            """Handle @mentions of the bot"""
            logger.info(f"Received mention: {event}")
            
            try:
                # Extract question from message
                text = event.get("text", "")
                # Remove bot mention
                question = text.split(">", 1)[-1].strip()
                
                # Get channel and thread info
                channel_id = event.get("channel")
                user_id = event.get("user")
                thread_ts = event.get("thread_ts") or event.get("ts")
                message_ts = event.get("ts")
                
                # Generate conversation ID
                # For thread mentions, use thread-specific ID to keep conversations separate
                if event.get("thread_ts"):
                    conversation_id = f"thread:{channel_id}:{thread_ts}"
                else:
                    conversation_id = f"{channel_id}:{user_id}"
                
                # If in a thread, get thread context
                if thread_ts:
                    thread_messages = self._get_thread_messages(channel_id, thread_ts)
                    # Index thread for context
                    rag_engine.index_slack_thread(thread_messages, channel_id)
                
                # Generate response using RAG with conversation memory
                response = rag_engine.generate_response(
                    question=question,
                    channel_id=channel_id,
                    conversation_id=conversation_id,
                    user_id=user_id,
                    message_ts=message_ts
                )
                
                # Store message in database
                self._store_message(event)
                
                # Reply in thread
                say(
                    text=response["answer"],
                    thread_ts=thread_ts
                )
                
            except Exception as e:
                logger.error(f"Error handling mention: {e}")
                say(
                    text=f"Sorry, I encountered an error: {str(e)}",
                    thread_ts=event.get("thread_ts") or event.get("ts")
                )
        
        @self.app.event("message")
        def handle_message(event, say):
            """Handle direct messages"""
            # Only respond to DMs, not channel messages
            if event.get("channel_type") == "im":
                logger.info(f"Received DM: {event}")
                
                try:
                    question = event.get("text", "")
                    user_id = event.get("user")
                    channel_id = event.get("channel")
                    message_ts = event.get("ts")
                    
                    # Generate conversation ID for DMs
                    conversation_id = f"dm:{user_id}"
                    
                    # Generate response with conversation memory
                    response = rag_engine.generate_response(
                        question=question,
                        channel_id=channel_id,
                        conversation_id=conversation_id,
                        user_id=user_id,
                        message_ts=message_ts
                    )
                    
                    # Store message
                    self._store_message(event)
                    
                    # Reply
                    say(text=response["answer"])
                    
                except Exception as e:
                    logger.error(f"Error handling message: {e}")
                    say(text=f"Sorry, I encountered an error: {str(e)}")
        
        @self.app.command("/wingman")
        def handle_command(ack, command, say):
            """Handle /wingman slash command"""
            ack()
            logger.info(f"Received command: {command}")
            
            try:
                question = command.get("text", "")
                
                if not question:
                    say(text="How can I help you? Please provide a question.")
                    return
                
                # Generate response
                response = rag_engine.generate_response(question)
                
                # Reply
                say(text=response["answer"])
                
            except Exception as e:
                logger.error(f"Error handling command: {e}")
                say(text=f"Sorry, I encountered an error: {str(e)}")
        
        @self.app.event("reaction_added")
        def handle_reaction(event):
            """Handle reactions to learn from user feedback"""
            logger.info(f"Reaction added: {event}")
            # Could use reactions like ✅ or ❌ to train the model
            pass
        
        @self.app.event("file_shared")
        def handle_file_shared(event, say):
            """Handle file uploads to index them automatically"""
            logger.info(f"File shared: {event}")
            
            try:
                file_id = event.get("file_id")
                user_id = event.get("user_id")
                channel_id = event.get("channel_id")
                
                # Get file info
                file_info = self.client.files_info(file=file_id)
                file_data = file_info.get("file", {})
                
                if not file_data:
                    logger.error(f"Could not retrieve file info for {file_id}")
                    return
                
                # Process and index the file
                success = self._process_and_index_file(file_data, channel_id, user_id)
                
                if success:
                    # React to the file to indicate it was indexed
                    if file_data.get("shares"):
                        # Get the message timestamp from shares
                        for channel, shares in file_data.get("shares", {}).get("public", {}).items():
                            if shares:
                                ts = shares[0].get("ts")
                                if ts:
                                    self.client.reactions_add(
                                        channel=channel,
                                        timestamp=ts,
                                        name="white_check_mark"
                                    )
                    
                    logger.info(f"Successfully indexed file: {file_data.get('name')}")
                else:
                    logger.warning(f"Could not index file: {file_data.get('name')}")
                    
            except Exception as e:
                logger.error(f"Error handling file upload: {e}")
    
    def _get_thread_messages(self, channel_id: str, thread_ts: str) -> list:
        """Retrieve all messages in a thread"""
        try:
            result = self.client.conversations_replies(
                channel=channel_id,
                ts=thread_ts
            )
            return result.get("messages", [])
        except Exception as e:
            logger.error(f"Error getting thread messages: {e}")
            return []
    
    def _store_message(self, event: Dict[str, Any]):
        """Store message in database"""
        try:
            db = SessionLocal()
            message = SlackMessage(
                message_ts=event.get("ts"),
                channel_id=event.get("channel"),
                user_id=event.get("user"),
                text=event.get("text"),
                thread_ts=event.get("thread_ts"),
                msg_metadata=event
            )
            db.add(message)
            db.commit()
            db.close()
        except Exception as e:
            logger.error(f"Error storing message: {e}")
    
    def _process_and_index_file(self, file_data: Dict[str, Any], channel_id: str, user_id: str) -> bool:
        """
        Download, process, and index a file from Slack
        
        Args:
            file_data: File metadata from Slack API
            channel_id: Channel where file was shared
            user_id: User who shared the file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            file_url = file_data.get("url_private")
            file_name = file_data.get("name", "unknown")
            file_type = file_data.get("filetype", "").lower()
            file_mimetype = file_data.get("mimetype", "")
            
            if not file_url:
                logger.error("File URL not available")
                return False
            
            # Download file content
            headers = {"Authorization": f"Bearer {settings.SLACK_BOT_TOKEN}"}
            with httpx.Client() as client:
                response = client.get(file_url, headers=headers)
            
            if response.status_code != 200:
                logger.error(f"Failed to download file: HTTP {response.status_code}")
                return False
            
            # Extract text content based on file type
            content = self._extract_file_content(response.content, file_type, file_mimetype)
            
            if not content:
                logger.warning(f"Could not extract content from {file_name}")
                return False
            
            # Store in database
            db = SessionLocal()
            document = Document(
                title=file_name,
                content=content,
                source="slack_upload",
                doc_metadata={
                    "channel_id": channel_id,
                    "user_id": user_id,
                    "file_id": file_data.get("id"),
                    "file_type": file_type,
                    "uploaded_at": file_data.get("created")
                }
            )
            db.add(document)
            db.commit()
            db.close()
            
            # Index in vector store
            rag_engine.index_document(
                title=file_name,
                content=content,
                source=f"slack_upload:{channel_id}"
            )
            
            logger.info(f"Successfully processed and indexed file: {file_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return False
    
    def _extract_file_content(self, file_bytes: bytes, file_type: str, mimetype: str) -> str:
        """
        Extract text content from various file types
        
        Args:
            file_bytes: Raw file content
            file_type: File extension
            mimetype: MIME type
            
        Returns:
            Extracted text content or empty string
        """
        try:
            # Plain text files
            if file_type in ["txt", "md", "markdown", "log", "csv", "json", "xml", "yaml", "yml"]:
                return file_bytes.decode("utf-8", errors="ignore")
            
            # Python files
            if file_type in ["py", "js", "ts", "java", "cpp", "c", "h", "go", "rb", "php", "sh"]:
                return file_bytes.decode("utf-8", errors="ignore")
            
            # PDF files
            if file_type == "pdf" or mimetype == "application/pdf":
                try:
                    from pypdf import PdfReader
                    pdf_file = io.BytesIO(file_bytes)
                    reader = PdfReader(pdf_file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
                except ImportError:
                    logger.warning("pypdf not available, cannot process PDF files")
                    return ""
            
            # Word documents
            if file_type in ["docx", "doc"] or mimetype in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    from docx import Document as DocxDocument
                    docx_file = io.BytesIO(file_bytes)
                    doc = DocxDocument(docx_file)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                except ImportError:
                    logger.warning("python-docx not available, cannot process Word documents")
                    return ""
            
            # HTML files
            if file_type in ["html", "htm"] or "text/html" in mimetype:
                # Simple HTML stripping (basic implementation)
                content = file_bytes.decode("utf-8", errors="ignore")
                # Remove script and style tags
                import re
                content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
                content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
                # Remove all HTML tags
                content = re.sub(r'<[^>]+>', '', content)
                # Clean up whitespace
                content = re.sub(r'\s+', ' ', content).strip()
                return content
            
            logger.warning(f"Unsupported file type: {file_type} ({mimetype})")
            return ""
            
        except Exception as e:
            logger.error(f"Error extracting file content: {e}")
            return ""
    
    def start(self):
        """Start the bot"""
        if settings.SLACK_APP_TOKEN:
            # Use Socket Mode for local development
            logger.info("Starting bot in Socket Mode...")
            handler = SocketModeHandler(self.app, settings.SLACK_APP_TOKEN)
            handler.start()
        else:
            logger.info("Socket Mode not enabled. Use Socket Mode for local development.")
            logger.info("For production, use a web server to handle events.")


# Global bot instance
slack_bot = SlackBot()
